from docx import Document
from docx.shared import Cm, Pt


class ReportGenerator:

    @staticmethod
    def generate(filename, session_id, records, points, image_path):

        d = Document()

        style = d.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12)

        d.add_heading(f"Отчёт по сессии №{session_id}", 1)
        d.add_paragraph(f"Количество кадров: {len(records)}")
        d.add_paragraph(f"Количество точек: {len(points)}")

        if points:
            temps = [float(p["temperature"])for p in points]
            d.add_paragraph(f"Мин. температура: {min(temps):.2f} °C")
            d.add_paragraph(f"Макс. температура: {max(temps):.2f} °C")
            d.add_paragraph(f"Средняя температура: "f"{sum(temps)/len(temps):.2f} °C")

        if image_path:
            try:
                d.add_picture(image_path, width=Cm(12))

            except Exception:
                pass

        d.save(filename)